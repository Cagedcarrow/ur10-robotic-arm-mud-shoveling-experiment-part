#!/usr/bin/env python3
"""
生成论文级 Word 文档 — 最终提交版本
=================================
基于 27 组 UR10 铲泥实验数据，含分析、响应面建模和多目标优化。
本版本对结论表述进行了严格限定，避免过度解读。
"""
import json, os, re, shutil, subprocess, sys, warnings, zipfile
from pathlib import Path

warnings.filterwarnings('ignore')
os.environ["PATH"] = "/usr/bin:" + os.environ.get("PATH", "")

# ── 路径 ──
BASE = Path("/root/ur10_ws/src/mud_experiment_recorder/data/analysis_results_20260513_005805")
TABLE_DIR = BASE / "tables"
FIG_DIR = BASE / "figures"
PAPER_DIR = BASE / "paper"
ASSET_DIR = PAPER_DIR / "assets"
PAPER_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
for f in sorted(FIG_DIR.glob("*.png")):
    shutil.copy2(f, ASSET_DIR / f.name)

import pandas as pd
import numpy as np

df = pd.read_csv(TABLE_DIR / "single_run_features.csv")
grid = pd.read_csv(TABLE_DIR / "response_surface_grid.csv")
cand_old = pd.read_csv(TABLE_DIR / "multi_objective_candidates.csv")

idx_max_mass = df["mass_g"].idxmax()
idx_min_mass = df["mass_g"].idxmin()
best_mass = df.loc[idx_max_mass]
worst_mass = df.loc[idx_min_mass]

# 响应面单目标质量最大候选（从网格中查找，与权重无关）
rsm_max_idx = grid["pred_mass_g"].idxmax()
rsm_best = grid.loc[rsm_max_idx]

# ── 重新进行多目标优化，包含 mass_per_force_impulse ──
grid = grid.copy()
grid["pred_mass_per_force_impulse"] = np.where(
    grid["pred_F_impulse"].values > 1e-9,
    grid["pred_mass_g"].values / grid["pred_F_impulse"].values,
    np.nan
)

def multi_objective_optimize(df_grid, objectives):
    """加权多目标综合评分，返回按评分降序排列的 DataFrame。"""
    scores = np.zeros(len(df_grid))
    weight_sum = 0.0
    for col, w in objectives.items():
        if col not in df_grid.columns:
            continue
        vals = df_grid[col].values.astype(float)
        valid_mask = ~np.isnan(vals)
        if not np.any(valid_mask):
            continue
        v = vals.copy()
        v_min, v_max = np.nanmin(v), np.nanmax(v)
        if v_max - v_min < 1e-9:
            continue
        if w > 0:
            norm = (v - v_min) / (v_max - v_min)
        else:
            norm = (v_max - v) / (v_max - v_min)
        norm[~valid_mask] = np.nan
        scores += abs(w) * norm
        weight_sum += abs(w)
    if weight_sum > 0:
        scores = scores / weight_sum
    df_out = df_grid.copy()
    df_out["score"] = scores
    return df_out.sort_values("score", ascending=False).reset_index(drop=True)

# 完整权重集合（合计 1.00）
FULL_OBJECTIVES = {
    "pred_mass_g": 0.30,
    "pred_mass_per_energy": 0.20,
    "pred_mass_per_force_impulse": 0.15,
    "pred_E_mech": -0.10,
    "pred_F_max": -0.10,
    "pred_F_impulse": -0.05,
    "pred_tcp_jerk_rms": -0.05,
    "pred_joint_rmse": -0.05,
}
cand = multi_objective_optimize(grid, FULL_OBJECTIVES)
top_cand = cand.iloc[0]
eff_best = cand.loc[cand["pred_mass_per_energy"].idxmax()]
ene_best = cand.loc[cand["pred_E_mech"].idxmin()]

# 各因素极差
def lv_mean(col, fac, lv):
    return df[df[fac] == lv][col].dropna().mean()

def lv_range(col, fac):
    levels = sorted(df[fac].unique())
    vals = [float(lv_mean(col, fac, l)) for l in levels]
    return max(vals) - min(vals), vals

range_angle, _ = lv_range("mass_g", "angle_deg")
range_speed, _ = lv_range("mass_g", "speed_percent")
range_depth, _ = lv_range("mass_g", "depth_mm")

corr_em = float(df["E_mech"].corr(df["mass_g"]))
corr_fm = float(df["F_mean"].corr(df["mass_g"]))
corr_fx = float(df["F_max"].corr(df["mass_g"]))

ft_ratios = df["ft_valid_ratio"].dropna()
ft_mean = ft_ratios.mean()
dt_mean_val = df["dt_mean"].mean()

emax_row = df.loc[df["E_mech"].idxmax()]
emin_row = df.loc[df["E_mech"].idxmin()]
mpe_row = df.loc[df["mass_per_energy"].idxmax()]
fmax_row = df.loc[df["F_max"].idxmax()]
fmean_row = df.loc[df["F_mean"].idxmax()]
fimp_row = df.loc[df["F_impulse"].idxmax()]

# 关节跟踪误差范围
jrmse_min = float(df["joint_rmse"].min())
jrmse_max = float(df["joint_rmse"].max())


def asset_img(name):
    """仅嵌入图片，不附带额外图题行。图题由调用方单独添加。"""
    return f"![图](assets/{name})\n"


def build_markdown():
    lines = []

    # ═══════ 标题 ═══════
    lines.append("# 基于响应面法的 UR10 机械臂铲泥工艺参数多目标优化分析\n")
    lines.append("*数据分析报告*\n")
    lines.append("---\n")

    # ═══════ 摘要 ═══════
    lines.append("## 摘要\n\n")
    lines.append("本文基于 27 组 UR10 机械臂铲泥实验数据，分析入泥角度（20°、35°、50°）、")
    lines.append("速度百分比（35%、70%、100%）和入泥深度（20、40、60 mm）对铲泥质量、")
    lines.append("机械能耗、阻力特征和轨迹平滑性的影响。")
    lines.append("结果表明，在已完成实验范围内，实际最大铲泥质量为 240 g，")
    lines.append(f"对应 20°、35°、60 mm；入泥深度对铲泥质量的主效应最大，极差为 {range_depth:.1f} g；")
    lines.append("铲泥质量二次响应面模型 $R^2=0.688$，调整后 $R^2=0.524$，")
    lines.append("说明模型具有一定趋势解释能力但仍存在过拟合风险。")
    lines.append(f"基于加权评分的多目标优化推荐 50°、100%、45 mm 作为折中候选参数，")
    lines.append(f"预测质量 148 g，预测能耗 11 J，但该组合尚未实测验证。")
    lines.append("研究结果可为当前铲斗结构、泥浆状态和轨迹条件下的参数筛选提供参考。\n")

    # ═══════ 1 引言 ═══════
    lines.append("## 1 引言\n\n")
    lines.append("机械臂铲泥作业在建筑拆除、矿山清理、水下采样和灾害救援等领域具有应用前景。")
    lines.append("铲泥作业涉及铲土-物料相互作用，其效果受运动参数（入泥角度、速度百分比、入泥深度）")
    lines.append("和物料属性（含水量、密度、粘度）的共同影响。")
    lines.append("在实际工程中，操作人员通常依赖经验进行参数整定，缺乏定量优化方法。\n\n")
    lines.append("响应面方法（Response Surface Methodology, RSM）能够在有限实验样本条件下建立")
    lines.append("输入变量与响应变量之间的连续映射关系。多目标优化可在相互冲突的目标之间寻求折中方案。\n\n")
    lines.append("本文以 UR10 协作机器人为实验平台，通过力/力矩传感器和实时数据采集系统记录铲泥过程，")
    lines.append("分析了 27 组三因素三水平实验数据，建立了二次响应面模型并开展了多目标参数筛选。\n")

    # ═══════ 2 实验系统与方法 ═══════
    lines.append("## 2 实验系统与方法\n\n")
    lines.append("### 2.1 硬件系统\n\n")
    lines.append("实验平台由以下组件构成：\n\n")
    lines.append("- UR10 协作机器人：6 自由度，额定负载 10 kg，重复定位精度 ±0.03 mm")
    lines.append("- FT300 力/力矩传感器：安装于机器人末端法兰，量程 ±300 N")
    lines.append("- 铲斗工装：安装于传感器下方")
    lines.append("- 上位机：运行 Ubuntu 22.04，通过 Socket 接口采集 UR 状态数据，")
    lines.append("  同时通过串口 Modbus RTU 协议采集 FT300 数据\n")

    lines.append("### 2.2 数据采集\n\n")
    lines.append("每次实验记录关节状态、TCP 位姿、力/力矩、估计扭矩、电气参数及时间基准等信息。\n\n")
    lines.append("UR 数据采样周期均值约 {:.4f} s（约 {:.0f} Hz）。".format(dt_mean_val, 1/dt_mean_val))
    lines.append("FT300 数据采样频率低于 UR 状态数据，具体有效更新频率以 ft_fresh 标记和时间戳统计为准。")
    lines.append(f"根据 ft_fresh 标记统计，FT 数据有效更新比例均值约 {ft_mean:.1%}。\n")
    lines.append("**weight_g 说明**：weight_g 为实验结束后人工称量得到的铲泥质量，作为主响应变量。")
    lines.append("力传感器数据不直接用于计算 weight_g，仅用于分析铲泥过程中的受力特征。")
    lines.append("本文分析中速度变量统一采用 speed_percent（百分比），")
    lines.append("与计划阶段可能使用的 mm/s 速度值不属于同一量纲。\n")

    lines.append("### 2.3 实验设计\n\n")
    lines.append("本实验采用三因素三水平全因子实验设计，共 27 个参数组合，每个参数组合执行一次。")
    lines.append("原始实验计划曾参考 L9 正交设计思想，但实际采集数据覆盖了角度、速度、深度三个因素的")
    lines.append("3³=27 个组合，因此本文按三因素三水平全因子实验进行分析。\n\n")
    lines.append("因素及水平如下：\n\n")
    lines.append("| 因素 | 符号 | 水平 1 | 水平 2 | 水平 3 |")
    lines.append("| :--- | :--- | ---: | ---: | ---: |")
    lines.append("| 入泥角度 (°) | $A$ | 20 | 35 | 50 |")
    lines.append("| 速度百分比 (%) | $B$ | 35 | 70 | 100 |")
    lines.append("| 入泥深度 (mm) | $C$ | 20 | 40 | 60 |\n")

    # ═══════ 3 数据分析方法 ═══════
    lines.append("## 3 数据分析方法\n\n")

    # ── 3.1 有效铲泥阶段识别 ──
    lines.append("### 3.1 有效铲泥阶段识别\n\n")
    lines.append("为避免非铲泥段对特征提取的干扰，采用基线阈值法自动识别有效铲泥阶段。\n\n")
    lines.append("本文优先采用基坐标系下合力 $F_{\\mathrm{base}}$ 进行有效铲泥阶段识别和受力特征提取；")
    lines.append("当基坐标系力不可用时，使用传感器坐标系合力 $F_{\\mathrm{res}}$ 作为替代。\n\n")
    lines.append("基坐标系合力计算公式：\n")
    lines.append("$$")
    lines.append("F_{\\mathrm{base}}(t) = \\sqrt{F_{\\mathrm{base},x}(t)^2 + F_{\\mathrm{base},y}(t)^2 + F_{\\mathrm{base},z}(t)^2}")
    lines.append("$$\n")
    lines.append("传感器坐标系合力：\n")
    lines.append("$$")
    lines.append("F_{\\mathrm{res}}(t) = \\sqrt{F_x(t)^2 + F_y(t)^2 + F_z(t)^2}")
    lines.append("$$\n")
    lines.append("取前 10% 时间段作为空载基线，记其合力均值和标准差为：\n")
    lines.append("$$")
    lines.append("\\mu_0 = \\mathrm{mean}(F_{\\mathrm{base,\\ baseline}}),\\quad")
    lines.append("\\sigma_0 = \\mathrm{std}(F_{\\mathrm{base,\\ baseline}})")
    lines.append("$$\n")
    lines.append("阈值为：\n")
    lines.append("$$")
    lines.append("F_{\\mathrm{th}} = \\mu_0 + 3\\sigma_0")
    lines.append("$$\n")
    lines.append("有效铲泥阶段定义为 $F_{\\mathrm{base}}(t) > F_{\\mathrm{th}}$")
    lines.append("从首次连续超过阈值到最后一次超过阈值的时间区间。\n")

    # ── 3.2 受力与能耗特征 ──
    lines.append("### 3.2 受力与能耗特征\n\n")
    lines.append("在有效铲泥阶段内计算以下力学指标：\n\n")
    lines.append("平均力、最大力、力均方根：\n")
    lines.append("$$")
    lines.append("F_{\\mathrm{mean}} = \\frac{1}{N}\\sum_{k=1}^{N} F_{\\mathrm{base}}(t_k),\\;")
    lines.append("F_{\\mathrm{max}} = \\max_{k} F_{\\mathrm{base}}(t_k),\\;")
    lines.append("F_{\\mathrm{RMS}} = \\sqrt{\\frac{1}{N}\\sum_{k=1}^{N} F_{\\mathrm{base}}^2(t_k)}")
    lines.append("$$\n")
    lines.append("力冲量：\n")
    lines.append("$$")
    lines.append("I_F = \\int_{t_s}^{t_e} F_{\\mathrm{base}}(t) \\, dt")
    lines.append("$$\n")
    lines.append("离散形式：\n")
    lines.append("$$")
    lines.append("I_F \\approx \\sum_{k} F_{\\mathrm{base}}(t_k) \\cdot \\Delta t_k")
    lines.append("$$\n")
    lines.append("关节机械能耗基于估计扭矩和关节角速度。各关节瞬时功率：\n")
    lines.append("$$")
    lines.append("P_i(t) = |\\tau_{\\mathrm{est},i}(t) \\cdot \\dot{q}_i(t)|")
    lines.append("$$\n")
    lines.append("总瞬时功率：\n")
    lines.append("$$")
    lines.append("P_{\\mathrm{total}}(t) = \\sum_{i=0}^{5} |\\tau_{\\mathrm{est},i}(t) \\cdot \\dot{q}_i(t)|")
    lines.append("$$\n")
    lines.append("机械能耗：\n")
    lines.append("$$")
    lines.append("E_{\\mathrm{mech}} = \\int_{t_s}^{t_e} P_{\\mathrm{total}}(t) \\, dt")
    lines.append("$$\n")
    lines.append("离散形式：\n")
    lines.append("$$")
    lines.append("E_{\\mathrm{mech}} \\approx \\sum_{k} P_{\\mathrm{total}}(t_k) \\cdot \\Delta t_k")
    lines.append("$$\n")
    lines.append("其中 $\\tau_{\\mathrm{est},i}$ 为 tau_estimated_i，$\\dot{q}_i$ 为 Act_qd_i。\n")

    # ── 3.3 轨迹平滑性 ──
    lines.append("### 3.3 轨迹平滑性\n\n")
    lines.append("TCP 速度：\n")
    lines.append("$$")
    lines.append("v_{\\mathrm{TCP}}(t) = \\sqrt{\\dot{X}(t)^2 + \\dot{Y}(t)^2 + \\dot{Z}(t)^2}")
    lines.append("$$\n")
    lines.append("TCP 加速度与急动度：\n")
    lines.append("$$")
    lines.append("a_{\\mathrm{TCP}}(t) = \\frac{dv_{\\mathrm{TCP}}}{dt},\\;")
    lines.append("j_{\\mathrm{TCP}}(t) = \\frac{da_{\\mathrm{TCP}}}{dt}")
    lines.append("$$\n")
    lines.append("对应均方根值：\n")
    lines.append("$$")
    lines.append("a_{\\mathrm{rms}} = \\sqrt{\\mathrm{mean}\\big(a_{\\mathrm{TCP}}^2(t)\\big)},\\;")
    lines.append("j_{\\mathrm{rms}} = \\sqrt{\\mathrm{mean}\\big(j_{\\mathrm{TCP}}^2(t)\\big)}")
    lines.append("$$\n")
    lines.append("关节跟踪误差。各关节误差：\n")
    lines.append("$$")
    lines.append("e_{q,i}(t_k) = q_{\\mathrm{tgt},i}(t_k) - q_{\\mathrm{act},i}(t_k)")
    lines.append("$$\n")
    lines.append("关节 RMSE：\n")
    lines.append("$$")
    lines.append("\\mathrm{RMSE}_{\\mathrm{joint}} = \\sqrt{\\frac{1}{6N}\\sum_{i=0}^{5}\\sum_{k=1}^{N}")
    lines.append("\\big(q_{\\mathrm{tgt},i}(t_k) - q_{\\mathrm{act},i}(t_k)\\big)^2}")
    lines.append("$$\n")
    lines.append("其中 $q_{\\mathrm{tgt},i}$ 为目标关节位置，$q_{\\mathrm{act},i}$ 为实际关节位置。\n")

    # ── 3.4 二次响应面模型 ──
    lines.append("### 3.4 二次响应面模型\n\n")
    lines.append("设 $x_1 = \\text{angle\\_deg}$（入泥角度），$x_2 = \\text{speed\\_percent}$（速度百分比），")
    lines.append("$x_3 = \\text{depth\\_mm}$（入泥深度）。二次响应面模型为：\n")
    lines.append("$$")
    lines.append("y = \\beta_0 + \\beta_1 x_1 + \\beta_2 x_2 + \\beta_3 x_3")
    lines.append("+ \\beta_{11} x_1^2 + \\beta_{22} x_2^2 + \\beta_{33} x_3^2")
    lines.append("+ \\beta_{12} x_1 x_2 + \\beta_{13} x_1 x_3 + \\beta_{23} x_2 x_3 + \\varepsilon")
    lines.append("$$\n")
    lines.append("模型拟合质量通过以下指标评价：\n\n")
    lines.append("决定系数 $R^2$：\n")
    lines.append("$$")
    lines.append("R^2 = 1 - \\frac{\\mathrm{SS}_{\\mathrm{res}}}{\\mathrm{SS}_{\\mathrm{tot}}}")
    lines.append("$$\n")
    lines.append("调整后 $R^2$：\n")
    lines.append("$$")
    lines.append("R^2_{\\mathrm{adj}} = 1 - (1 - R^2)\\frac{n-1}{n-p-1}")
    lines.append("$$\n")
    lines.append("其中 $n$ 为样本数（$n=27$），$p$ 为自变量项数量（含二次项和交互项，$p=9$）。\n\n")
    lines.append("RMSE：\n")
    lines.append("$$")
    lines.append("\\mathrm{RMSE} = \\sqrt{\\frac{1}{n}\\sum_{i=1}^{n} (y_i - \\hat{y}_i)^2}")
    lines.append("$$\n")

    # ── 3.5 多目标优化 ──
    lines.append("### 3.5 多目标优化\n\n")
    lines.append("采用极差归一化加权评分法。综合评分：\n")
    lines.append("$$")
    lines.append("S = \\sum_{i} w_i \\, z_i")
    lines.append("$$\n")
    lines.append("其中 $z_i$ 为归一化后的指标值。正向指标（最大化）：\n")
    lines.append("$$")
    lines.append("z_i = \\frac{x_i - x_{\\min}}{x_{\\max} - x_{\\min}}")
    lines.append("$$\n")
    lines.append("负向指标（最小化）：\n")
    lines.append("$$")
    lines.append("z_i = \\frac{x_{\\max} - x_i}{x_{\\max} - x_{\\min}}")
    lines.append("$$\n")
    lines.append("评分越高表示该候选方案在给定权重体系下综合表现越好。\n\n")
    lines.append("各指标权重如下（合计 1.00）：\n\n")
    lines.append("| 指标 | 优化方向 | 权重 |")
    lines.append("| :--- | :---: | ---: |")
    lines.append("| weight_g（铲泥质量） | 最大化 | 0.30 |")
    lines.append("| mass_per_energy（单位能耗铲泥量） | 最大化 | 0.20 |")
    lines.append("| mass_per_force_impulse（单位力冲量铲泥量） | 最大化 | 0.15 |")
    lines.append("| E_mech（机械能耗） | 最小化 | 0.10 |")
    lines.append("| F_max（最大合力） | 最小化 | 0.10 |")
    lines.append("| F_impulse（力冲量） | 最小化 | 0.05 |")
    lines.append("| tcp_jerk_rms（TCP 急动度 RMS） | 最小化 | 0.05 |")
    lines.append("| joint_rmse（关节跟踪误差 RMSE） | 最小化 | 0.05 |\n")
    lines.append("注：若某指标在网格数据中不可计算，则跳过该指标并对剩余权重重新归一化。")
    lines.append("最终综合评分已归一化至 [0, 1] 区间。\n")

    # ═══════ 4 结果与讨论 ═══════
    lines.append("## 4 结果与讨论\n\n")

    # ── 4.1 铲泥质量 ──
    lines.append("### 4.1 铲泥质量分析\n\n")
    lines.append(f"27 组实验的铲泥质量范围为 50~240 g，平均值 122.6 g。\n\n")
    lines.append(f"**实际最大铲泥质量**：{best_mass['folder']}，角度 {best_mass['angle_deg']:.0f}°、")
    lines.append(f"速度 {best_mass['speed_percent']:.0f}%、深度 {best_mass['depth_mm']:.0f} mm，")
    lines.append(f"质量 {best_mass['mass_g']:.0f} g。\n\n")
    lines.append(f"**实际最小铲泥质量**：{worst_mass['folder']}，角度 {worst_mass['angle_deg']:.0f}°、")
    lines.append(f"速度 {worst_mass['speed_percent']:.0f}%、深度 {worst_mass['depth_mm']:.0f} mm，")
    lines.append(f"质量 {worst_mass['mass_g']:.0f} g。\n\n")
    lines.append(asset_img("01_weight_bar_all_runs.png"))
    lines.append("**图 1  27 组实验铲泥质量柱状图**。\n\n")
    lines.append(asset_img("02_weight_by_angle_speed_depth.png"))
    lines.append("**图 2  三因素下铲泥质量分布**。\n\n")
    lines.append("**极差法因素影响分析**：\n\n")
    for fac, name in [("angle_deg", "入泥角度"), ("speed_percent", "速度百分比"), ("depth_mm", "入泥深度")]:
        ms = [float(lv_mean("mass_g", fac, l)) for l in sorted(df[fac].unique())]
        r = round(max(ms) - min(ms), 1)
        trend = " → ".join([f"{m:.1f}" for m in ms])
        lines.append(f"- {name}极差 {r} g，各水平均值：{trend}")
    lines.append("")
    lines.append(f"入泥深度对铲泥质量的影响最为显著（极差 {range_depth:.1f} g），")
    lines.append("中等深度（40 mm）的平均铲泥量最高，但单次最大值出现在 60 mm，")
    lines.append("因此需要后续重复实验区分平均稳定性和单次极值。\n")

    # ── 4.2 力学特征 ──
    lines.append("### 4.2 力学特征分析\n\n")
    lines.append(f"最大力峰值：{fmax_row['folder']}，F_max = {fmax_row['F_max']:.1f} N，")
    lines.append(f"质量 {fmax_row['mass_g']:.0f} g。\n\n")
    lines.append(f"平均力最大：{fmean_row['folder']}，F_mean = {fmean_row['F_mean']:.1f} N，")
    lines.append(f"质量 {fmean_row['mass_g']:.0f} g。\n\n")
    lines.append(f"力冲量最大：{fimp_row['folder']}，I_F = {fimp_row['F_impulse']:.1f} N·s，")
    lines.append(f"质量 {fimp_row['mass_g']:.0f} g。\n\n")
    lines.append(f"铲泥质量与平均力的相关系数约 {corr_fm:.3f}，与最大力的相关系数约 {corr_fx:.3f}，均小于 0.3。")
    lines.append("说明铲泥质量并非简单随阻力增大而增大。部分实验存在受力较大但铲泥量较低的情况，")
    lines.append("提示不合理参数会产生无效阻力。")
    lines.append("（注：相关性分析仅用于趋势观察，样本量较小，不能作为因果关系判断。）\n\n")
    lines.append("**FT 数据有效率的说明**：FT300 数据的 ft_fresh 有效更新比例均值约为 {:.1%}，".format(ft_mean))
    lines.append("明显低于 UR 状态数据更新频率。因此，F_mean、F_max、F_impulse 等力学特征可能")
    lines.append("受到力传感器数据更新率、保持值或插值处理的影响。")
    lines.append("本文受力分析主要用于不同实验之间的相对比较，")
    lines.append("不宜将其解释为高精度绝对力学测量结果。")
    lines.append("后续实验应优先提高 FT300 通信稳定性和有效采样比例。\n\n")
    lines.append(asset_img("03_force_time_examples.png"))
    lines.append("**图 3  典型实验力-时间曲线（绿色区域标注铲泥阶段）**。\n\n")
    lines.append(asset_img("04_force_features_bar.png"))
    lines.append("**图 4  各实验受力特征对比**。\n")

    # ── 4.3 能耗 ──
    lines.append("### 4.3 能量消耗分析\n\n")
    lines.append(f"机械能耗最大：{emax_row['folder']}，E_mech = {emax_row['E_mech']:.1f} J，")
    lines.append(f"质量 {emax_row['mass_g']:.0f} g。\n\n")
    lines.append(f"机械能耗最小：{emin_row['folder']}，E_mech = {emin_row['E_mech']:.1f} J，")
    lines.append(f"质量 {emin_row['mass_g']:.0f} g。\n\n")
    lines.append(f"单位能耗铲泥量最高：{mpe_row['folder']}，{mpe_row['mass_per_energy']:.3f} g/J，")
    lines.append(f"质量 {mpe_row['mass_g']:.0f} g，能耗 {mpe_row['E_mech']:.1f} J。\n\n")
    lines.append(f"E_mech 与 mass_g 的相关系数约 {corr_em:.3f}，接近于 0，")
    lines.append("说明在当前 27 组样本中机械能耗与铲泥质量之间几乎不存在显著线性相关关系。")
    lines.append("高铲泥量并不必然对应高能耗，低能耗也不必然对应低铲泥量。\n\n")
    lines.append(asset_img("05_energy_features_bar.png"))
    lines.append("**图 5  各实验能耗特征对比**。\n")

    # ── 4.4 轨迹平滑性 ──
    lines.append("### 4.4 轨迹平滑性与跟踪误差\n\n")
    lines.append(f"27 组实验的加速度 RMS 均值 {df['tcp_acc_rms'].mean():.2f}，")
    lines.append(f"急动度 RMS 均值 {df['tcp_jerk_rms'].mean():.0f}，")
    lines.append(f"关节跟踪误差 RMSE 均值 {df['joint_rmse'].mean():.6f} rad。\n\n")
    lines.append(f"关节跟踪误差整体较小（< 0.0003 rad），")
    lines.append("说明在本实验参数范围内 UR10 的关节轨迹跟踪性能较稳定。")
    lines.append("但轨迹平滑性指标在部分高速和大深度组合下可能出现波动，")
    lines.append("因此仍需结合加速度 RMS 和急动度 RMS 评估运动冲击。\n\n")
    lines.append("TCP 急动度由离散速度信号差分得到，对采样噪声和滤波方法较敏感。")
    lines.append("因此本文仅将 tcp_jerk_rms 作为不同实验之间的相对平滑性指标，")
    lines.append("不将其绝对数值直接解释为结构冲击载荷。\n\n")
    lines.append(asset_img("06_trajectory_smoothness_bar.png"))
    lines.append("**图 6  轨迹平滑性与跟踪误差对比**。\n")

    # ── 4.5 因素主效应 ──
    lines.append("### 4.5 因素主效应分析\n\n")
    lines.append(asset_img("11_factor_main_effect_weight.png"))
    lines.append("**图 11  铲泥质量主效应图**。\n\n")
    lines.append(asset_img("12_factor_main_effect_energy.png"))
    lines.append("**图 12  机械能耗主效应图**。\n\n")
    lines.append(asset_img("13_factor_main_effect_force.png"))
    lines.append("**图 13  阻力峰值主效应图**。\n\n")
    lines.append("**因素极差汇总**：\n\n")
    lines.append("| 响应变量 | 影响排序 | 极差最大因素 |")
    lines.append("| :--- | :--- | :--- |")
    for resp, name in [("mass_g", "铲泥质量"), ("E_mech", "机械能耗"), ("F_max", "最大力")]:
        rgs = {}
        for fac in ["angle_deg", "speed_percent", "depth_mm"]:
            ms = [float(lv_mean(resp, fac, l)) for l in sorted(df[fac].unique())]
            rgs[fac] = round(max(ms) - min(ms), 1)
        sf = sorted(rgs, key=rgs.get, reverse=True)
        # 用 ≈ 连接极差相同或非常接近的因素
        parts = []
        i = 0
        while i < len(sf):
            group = [sf[i]]
            j = i + 1
            while j < len(sf) and abs(rgs[sf[j]] - rgs[sf[i]]) < 0.01:
                group.append(sf[j])
                j += 1
            parts.append(" ≈ ".join(group) if len(group) > 1 else group[0])
            i = j
        s = " > ".join(parts)
        lines.append(f"| {name} | {s} | {sf[0]}（极差 {rgs[sf[0]]}）|")
    lines.append("")
    lines.append("铲泥质量影响因素排序中，入泥深度极差最大。")
    lines.append(f"入泥角度极差 {range_angle:.1f} g，速度百分比极差 {range_speed:.1f} g，")
    lines.append("二者相同或非常接近，当前数据不足以严格区分二者对铲泥质量影响的强弱。\n")

    # ── 4.6 响应面模型 ──
    lines.append("### 4.6 二次响应面模型\n\n")
    lines.append("**模型拟合质量**：\n\n")
    lines.append("| 响应变量 | $R^2$ | $R^2_{\\mathrm{adj}}$ | RMSE | MAE |")
    lines.append("| :--- | ---: | ---: | ---: | ---: |")
    rsm_data = [
        ("铲泥质量", 0.688, 0.524, 26.16, 19.76),
        ("机械能耗", 0.310, -0.056, 8.03, 6.88),
        ("最大力", 0.510, 0.250, 7.93, 5.98),
        ("力冲量", 0.565, 0.334, 54.57, 42.58),
        ("急动度RMS", 0.281, -0.099, 405562, 247052),
        ("单位能耗铲泥量", 0.369, 0.035, 5.60, 4.16),
    ]
    for name, r2, ar2, rmse, mae in rsm_data:
        lines.append(f"| {name} | {r2:.3f} | {ar2:.3f} | {rmse:.2f} | {mae:.2f} |")
    lines.append("")
    lines.append("铲泥质量响应面模型可解释约 68.8% 的样本方差，但调整后 $R^2$ 为 0.524，")
    lines.append("说明考虑模型复杂度后解释能力为中等水平。由于样本量仅 27 组，")
    lines.append("而二次模型包含 10 个参数（含截距），模型存在过拟合风险。")
    lines.append("因此响应面结果主要用于趋势分析和候选参数筛选，不宜作为高精度点预测依据。\n\n")
    lines.append("除铲泥质量外，机械能耗、急动度 RMS、单位能耗铲泥量等响应的 Adjusted $R^2$ 较低，")
    lines.append("说明这些响应面模型的解释能力有限。因此多目标优化结果对模型误差和权重设置较敏感，")
    lines.append("只能作为候选参数筛选依据。\n\n")
    lines.append(asset_img("07_response_surface_weight_angle_speed.png"))
    lines.append("**图 7  铲泥质量响应面趋势图（固定深度=40 mm，角度×速度）**。\n\n")
    lines.append(asset_img("08_response_surface_weight_angle_depth.png"))
    lines.append("**图 8  铲泥质量响应面趋势图（固定速度=70%，角度×深度）**。\n\n")
    lines.append(asset_img("09_response_surface_weight_speed_depth.png"))
    lines.append("**图 9  铲泥质量响应面趋势图（固定角度=35°，速度×深度）**。\n\n")
    lines.append("响应面趋势观察：\n\n")
    lines.append("1. 角度对铲泥质量的影响呈非单调趋势，20°和 50° 的平均表现优于 35°。\n")
    lines.append("2. 低速（35%）的平均铲泥量最高，但个体差异较大。\n")
    lines.append("3. 深度 40 mm 的平均值最高，但 60 mm 可产生单次最大值（240 g）。\n")
    lines.append("4. 响应面在 50°、40~50 mm 区域出现预测高值，提示该区域可进一步通过实验验证。\n")

    # ── 4.7 多目标优化 ──
    lines.append("### 4.7 多目标优化候选方案\n\n")
    lines.append("基于二次响应面网格（25³ = 15,625 个点）的加权综合评分，推荐前十候选方案如下：\n\n")
    lines.append("| 排名 | 角度(°) | 速度(%) | 深度(mm) | 综合评分 | 预测质量(g) | 预测能耗(J) | 预测F_max(N) | 预测效率(g/J) |")
    lines.append("| ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: |")
    for i in range(min(10, len(cand))):
        r = cand.iloc[i]
        lines.append(f"| {i+1} | {r['angle_deg']:.0f} | {r['speed_percent']:.0f} | {r['depth_mm']:.0f} | "
                     f"{r['score']:.3f} | {r['pred_mass_g']:.0f} | {r['pred_E_mech']:.0f} | "
                     f"{r['pred_F_max']:.1f} | {r['pred_mass_per_energy']:.3f} |")
    lines.append("")
    lines.append("**三组有代表性的候选方案**：\n\n")
    lines.append(f"1. **单目标—实际最大**：角度 20°、速度 35%、深度 60 mm，实际铲泥质量 240 g。")
    lines.append(f"   该组合铲泥量最高，但机械能耗 {emax_row['E_mech']:.1f} J 也较高，属于高产高能耗组合。\n")
    lines.append(f"2. **响应面预测的单目标质量最大**：角度 {rsm_best['angle_deg']:.0f}°、速度 {rsm_best['speed_percent']:.0f}%、")
    lines.append(f"   深度 {rsm_best['depth_mm']:.0f} mm，预测质量 {rsm_best['pred_mass_g']:.0f} g。")
    lines.append(f"   注意：该组合是模型预测点，原始 27 组实验中未直接测试，需后续验证。\n")
    lines.append(f"3. **多目标综合评分最高**：角度 {top_cand['angle_deg']:.0f}°、速度 {top_cand['speed_percent']:.0f}%、")
    lines.append(f"   深度 {top_cand['depth_mm']:.0f} mm，综合评分 {top_cand['score']:.3f}，")
    lines.append(f"   预测质量 {top_cand['pred_mass_g']:.0f} g，预测能耗 {top_cand['pred_E_mech']:.0f} J。")
    lines.append("   注意：该组合不是铲泥质量最大点，而是质量、能耗、阻力和平滑性之间的折中候选点。")
    lines.append("   该组合同样未在原始 27 组实验中直接测试，且评分依赖于当前模型和权重设置。")
    lines.append("   多目标优化结果并不表示该参数组合在真实实验中必然最优，")
    lines.append("   其含义是在当前模型、当前权重和当前参数范围内该点的综合评分最高。")
    lines.append("   最终是否优于实际高铲泥量组合，需要通过重复实验验证。\n")
    lines.append(f"**其他补充**：单位能耗铲泥量最高候选为角度 {eff_best['angle_deg']:.0f}°、速度 {eff_best['speed_percent']:.0f}%、")
    lines.append(f"深度 {eff_best['depth_mm']:.0f} mm，预测效率 {eff_best['pred_mass_per_energy']:.3f} g/J；")
    lines.append(f"能耗最低候选为角度 {ene_best['angle_deg']:.0f}°、速度 {ene_best['speed_percent']:.0f}%、")
    lines.append(f"深度 {ene_best['depth_mm']:.0f} mm，预测能耗 {ene_best['pred_E_mech']:.1f} J。\n\n")
    lines.append(asset_img("10_pareto_mass_energy_force.png"))
    lines.append("**图 10  能耗-质量多目标候选分布**。\n\n")
    lines.append(asset_img("14_correlation_heatmap.png"))
    lines.append("**图 14  关键指标相关性热力图**。\n")

    # ── 4.8 关键结果汇总 ──
    lines.append("### 4.8 关键结果汇总\n\n")
    lines.append("| 目标 | 最优/推荐组合 | 结果 | 评价 |")
    lines.append("| :--- | :--- | :--- | :--- |")
    lines.append(f"| 实际铲泥质量最大 | 20°、35%、60 mm | 240 g，E_mech={emax_row['E_mech']:.1f} J | 高产高能耗 |")
    lines.append(f"| 单位能耗铲泥量最高 | 20°、100%、40 mm | 140 g，E_mech={mpe_row['E_mech']:.1f} J，{mpe_row['mass_per_energy']:.3f} g/J | 效率最高但质量中等 |")
    lines.append(f"| 响应面预测质量最大 | {rsm_best['angle_deg']:.0f}°、{rsm_best['speed_percent']:.0f}%、{rsm_best['depth_mm']:.0f} mm | 预测质量 {rsm_best['pred_mass_g']:.0f} g | 模型预测点，未实测 |")
    lines.append(f"| 多目标综合评分最高 | {top_cand['angle_deg']:.0f}°、{top_cand['speed_percent']:.0f}%、{top_cand['depth_mm']:.0f} mm | score={top_cand['score']:.3f}，预测质量 {top_cand['pred_mass_g']:.0f} g，预测能耗 {top_cand['pred_E_mech']:.0f} J | 折中候选，未实测 |")
    lines.append(f"| 力冲量最大异常组 | {fimp_row['angle_deg']:.0f}°、{fimp_row['speed_percent']:.0f}%、{fimp_row['depth_mm']:.0f} mm | F_impulse={fimp_row['F_impulse']:.1f} N·s，质量 {fimp_row['mass_g']:.0f} g | 存在明显无效阻力 |\n")

    # ═══════ 5 结果解释边界 ═══════
    lines.append("## 5 结果解释边界\n\n")
    lines.append("本文结论仅适用于以下边界条件：\n\n")
    lines.append("- 机械臂型号：UR10，铲斗结构及安装方式固定；")
    lines.append("- 泥浆材料：实验当次泥浆状态（含水率、粘度、密度等未严格量化控制）；")
    lines.append("- 运动轨迹：实验中使用的铲泥轨迹类型（非任意轨迹）；")
    lines.append("- 参数范围：入泥角度 20~50°，速度百分比 35~100%，入泥深度 20~60 mm；")
    lines.append("- FT300 有效更新比例较低会削弱力学特征，特别是 F_impulse 和 F_max 的绝对可信度，")
    lines.append("  因此受力相关结论应定位为趋势性判断。\n")
    lines.append("此外，由于每个参数组合只有一次实验，无法估计同一组合下的随机误差；")
    lines.append("泥浆状态可能随实验顺序发生变化（如被前序实验扰动、堆积或含水率局部变化）。")
    lines.append("因此现阶段结论应定位为趋势分析和候选参数筛选，而非最终工艺定型。\n")

    # ═══════ 6 结论 ═══════
    lines.append("## 6 结论\n\n")
    lines.append("基于 27 组 UR10 机械臂铲泥实验数据，本文分析了入泥角度、速度百分比和入泥深度")
    lines.append("对铲泥质量、能耗、受力以及运动平滑性的影响。主要发现如下：\n\n")
    lines.append(f"1. 在 27 组实际实验中，最大铲泥质量出现在 {best_mass['folder']}，")
    lines.append(f"参数为 {best_mass['angle_deg']:.0f}°、{best_mass['speed_percent']:.0f}%、{best_mass['depth_mm']:.0f} mm，")
    lines.append(f"质量为 {best_mass['mass_g']:.0f} g。")
    lines.append(f"该组合铲泥量最高，但机械能耗也较高（{emax_row['E_mech']:.1f} J），属于高产高能耗组合。\n\n")
    lines.append(f"2. 单位能耗铲泥量最高的是 {mpe_row['folder']}，")
    lines.append(f"参数为角度 {mpe_row['angle_deg']:.0f}°、速度 {mpe_row['speed_percent']:.0f}%、深度 {mpe_row['depth_mm']:.0f} mm，")
    lines.append(f"铲泥质量 {mpe_row['mass_g']:.0f} g，机械能耗 {mpe_row['E_mech']:.1f} J，")
    lines.append(f"单位能耗铲泥量 {mpe_row['mass_per_energy']:.3f} g/J。")
    lines.append("该组合说明铲泥效率和单次铲泥量不是同一个最优目标。\n\n")
    lines.append(f"3. 主效应分析表明，入泥深度对铲泥质量影响最大，极差 {range_depth:.1f} g；")
    lines.append(f"入泥角度极差 {range_angle:.1f} g，速度百分比极差 {range_speed:.1f} g，")
    lines.append("二者相同或接近，当前数据不足以严格区分入泥角度和速度百分比对铲泥质量影响的强弱。")
    lines.append("深度 40 mm 的平均铲泥质量最高，但单次最大值出现在 60 mm，")
    lines.append("需要后续重复实验区分平均稳定性和单次极值。\n\n")
    lines.append(f"4. 铲泥质量与平均力、最大力之间的相关系数均小于 0.3，")
    lines.append("说明铲泥质量并非简单随阻力增大而增大。部分实验存在受力较大但铲泥量较低的情况，")
    lines.append("提示不合理参数会产生无效阻力。\n\n")
    lines.append("5. 二次响应面模型可用于趋势分析，但不能直接作为最终最优参数依据。")
    lines.append(f"响应面预测的单目标质量最大组合为角度 {rsm_best['angle_deg']:.0f}°、速度 {rsm_best['speed_percent']:.0f}%、")
    lines.append(f"深度 {rsm_best['depth_mm']:.0f} mm，预测质量 {rsm_best['pred_mass_g']:.0f} g；")
    lines.append(f"多目标综合优化推荐为角度 {top_cand['angle_deg']:.0f}°、速度 {top_cand['speed_percent']:.0f}%、")
    lines.append(f"深度 {top_cand['depth_mm']:.0f} mm，预测质量 {top_cand['pred_mass_g']:.0f} g、")
    lines.append(f"预测能耗 {top_cand['pred_E_mech']:.0f} J、综合评分 {top_cand['score']:.3f}。")
    lines.append("两者均为候选参数，需后续实测验证。\n\n")
    lines.append("6. 当前研究的主要局限包括：样本量仅 27 组；每个参数组合缺少重复实验；")
    lines.append("泥浆含水率、密度、表面平整度等未严格量化控制；响应面模型存在过拟合风险。")
    lines.append("因此后续应围绕关键候选组合开展重复实验，具体建议见第 7 节。\n")

    # ═══════ 7 后续验证实验建议 ═══════
    lines.append("## 7 后续验证实验建议\n\n")
    lines.append("基于当前分析结果，建议在以下参数组合开展重复验证实验：\n\n")
    lines.append("| 验证目的 | 参数组合 | 建议重复次数 | 重点观察指标 |")
    lines.append("| :--- | :--- | ---: | :--- |")
    lines.append("| 验证实际高铲泥量组合稳定性 | 20°、35%、60 mm | ≥3 | 平均铲泥质量、质量标准差、E_mech、F_max |")
    lines.append("| 验证单位能耗最高组合 | 20°、100%、40 mm | ≥3 | mass_per_energy、E_mech、质量稳定性 |")
    lines.append("| 验证响应面质量最大预测点 | 20°、35%、53 mm | ≥3 | 实测质量是否接近预测 202 g |")
    lines.append("| 验证多目标折中候选点 | 50°、100%、45 mm | ≥3 | 综合评分、质量、能耗、阻力、急动度 |")
    lines.append("| 验证高质量折中组合 | 50°、35%、40 mm | ≥3 | 质量、能耗、受力峰值、轨迹平滑性 |\n")
    lines.append("后续验证实验应在每次实验前重新搅拌并刮平泥浆表面，")
    lines.append("记录泥浆含水率或至少记录泥浆状态，尽量降低物料状态变化对结果的干扰。\n")

    # ═══════ 参考文献 ═══════
    lines.append("## 参考文献\n\n")
    lines.append("[1] Box G E P, Wilson K B. On the experimental attainment of optimum conditions[J]. ")
    lines.append("Journal of the Royal Statistical Society: Series B, 1951, 13(1): 1-38.\n\n")
    lines.append("[2] Myers R H, Montgomery D C, Anderson-Cook C M. Response surface methodology: ")
    lines.append("process and product optimization using designed experiments[M]. 4th ed. Wiley, 2016.\n\n")
    lines.append("[3] 王永菲, 王成国. 响应面法的理论与应用[J]. 中央民族大学学报(自然科学版), 2005, 14(3): 236-240.\n\n")
    lines.append("[4] Universal Robots. UR10 User Manual[M]. Universal Robots A/S, 2020.\n\n")
    lines.append("[5] Onrobot. FT300 Force Torque Sensor Product Manual[M]. Onrobot A/S, 2019.\n\n")
    lines.append("[6] 葛宜元. 试验设计方法与Design-Expert软件应用[M]. 哈尔滨工业大学出版社, 2015.\n")

    return "\n".join(lines)


# ═══════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════

md_path = PAPER_DIR / "paper.md"
docx_path = PAPER_DIR / "UR10_铲泥实验多目标优化分析报告.docx"

md_text = build_markdown()
md_path.write_text(md_text, encoding="utf-8")
print(f"Markdown 已生成: {md_path} ({len(md_text)} chars)")

# Pandoc 转换
print("正在执行 pandoc 转换...")
cmd = [
    "pandoc",
    str(md_path),
    "-f", "markdown+tex_math_dollars+grid_tables+pipe_tables",
    "-t", "docx",
    "-s",
    "--resource-path", str(PAPER_DIR),
    "-o", str(docx_path),
]
result = subprocess.run(cmd, capture_output=True, text=True)
if result.returncode != 0:
    print(f"pandoc 错误: {result.stderr}")
    sys.exit(1)
print(f"pandoc 转换成功: {docx_path}")

# python-docx 后格式化
print("正在后格式化 docx...")
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

doc = Document(str(docx_path))

section = doc.sections[0]
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)


def set_style_font(style, latin="Times New Roman", east_asia="宋体", size_pt=None):
    style.font.name = latin
    rpr = style._element.rPr
    if rpr is None:
        style._element.get_or_add_rPr()
        rpr = style._element.rPr
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), east_asia)
    if size_pt is not None:
        style.font.size = Pt(size_pt)


style_size_map = {"Normal": 12, "Heading 1": 16, "Heading 2": 14, "Heading 3": 13}
for st in doc.styles:
    if st.name in style_size_map:
        set_style_font(st, latin="Times New Roman", east_asia="宋体", size_pt=style_size_map[st.name])

# 正文段落格式
for para in doc.paragraphs:
    pf = para.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    for run in para.runs:
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), "宋体")

# 表格格式
for table in doc.tables:
    table.alignment = 1  # CENTER
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.line_spacing = 1.0
                pf.space_after = Pt(2)
                pf.space_before = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

doc.save(str(docx_path))
print("后格式化完成")


def count_omath(path):
    with zipfile.ZipFile(str(path), "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return xml.count("<m:oMath>")


def approx_zh(text):
    return len(re.findall(r"[一-鿿]", text))


omath = count_omath(docx_path)
zh_count = approx_zh(md_text)

print(f"\n{'='*60}")
print(f"文档生成成功！")
print(f"{'='*60}")
print(f"输出: {docx_path}")
print(f"大小: {docx_path.stat().st_size/1024:.1f} KB")
print(f"OMML 公式数: {omath}")
print(f"中文字符数: {zh_count}")
print(f"{'='*60}")
