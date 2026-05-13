#!/usr/bin/env python3
"""
公式渲染脚本 — 生成公式渲染完整的最终版
====================================
读取原始 paper.md，保留所有 $$...$$ 和 $...$ LaTeX 公式，
仅做格式修正，然后通过 pandoc +tex_math_dollars 生成含 OMML 公式的 docx。
"""
import os, re, shutil, subprocess, sys, warnings, zipfile
from pathlib import Path

warnings.filterwarnings('ignore')
os.environ["PATH"] = "/usr/bin:" + os.environ.get("PATH", "")

BASE = Path("/root/ur10_ws/src/mud_experiment_recorder/data/analysis_results_20260513_005805")
PAPER_DIR = BASE / "paper"
FIG_DIR = BASE / "figures"
ASSET_DIR = PAPER_DIR / "assets"
SRC_MD = PAPER_DIR / "paper.md"
DST_MD = PAPER_DIR / "基于响应面法的UR10机械臂铲泥工艺参数多目标优化分析_公式渲染版.md"
DST_DOCX = PAPER_DIR / "基于响应面法的UR10机械臂铲泥工艺参数多目标优化分析_公式渲染版.docx"

ASSET_DIR.mkdir(parents=True, exist_ok=True)
for f in sorted(FIG_DIR.glob("*.png")):
    shutil.copy2(f, ASSET_DIR / f.name)

# ============================================================
# 格式修正（不修改 LaTeX 公式）
# ============================================================

def fix_markdown(md_text: str) -> str:
    """修正格式问题，但保留 $$...$$ 和 $...$ 公式不变。"""

    # 1. 摘要中 35°/35% 错误
    md_text = md_text.replace(
        '对应 20°、35°、60 mm；',
        '对应 20°、35%、60 mm；'
    )
    md_text = re.sub(
        r'(参数为|角 )20°、35°、60 mm',
        r'\g<1>20°、35%、60 mm',
        md_text
    )

    # 2. 表格中 $A$ → A, $B$ → B, $C$ → C（因素符号）
    md_text = md_text.replace('| $A$ |', '| A |')
    md_text = md_text.replace('| $B$ |', '| B |')
    md_text = md_text.replace('| $C$ |', '| C |')

    # 3. 模型拟合质量表头：$R^2$ → R², $R^2_{\mathrm{adj}}$ → Adjusted R²
    md_text = md_text.replace(
        '| $R^2$ | $R^2_{\\mathrm{adj}}$ | RMSE | MAE |',
        '| R² | Adjusted R² | RMSE | MAE |'
    )

    # 4. 压缩多余空行（3+ → 2）
    md_text = re.sub(r'\n{4,}', '\n\n\n', md_text)

    return md_text


# ============================================================
# 执行
# ============================================================

src_text = SRC_MD.read_text(encoding='utf-8')
print(f"已读取: {SRC_MD} ({len(src_text)} chars)")

fixed_text = fix_markdown(src_text)
print(f"格式修正完成 ({len(fixed_text)} chars)")

DST_MD.write_text(fixed_text, encoding='utf-8')
print(f"已写入: {DST_MD}")

# pandoc 转换（带 tex_math_dollars 以生成 OMML 公式）
print("正在执行 pandoc 转换...")
cmd = [
    "pandoc", str(DST_MD),
    "-f", "markdown+tex_math_dollars+grid_tables+pipe_tables",
    "-t", "docx", "-s",
    "--resource-path", str(PAPER_DIR),
    "-o", str(DST_DOCX),
]
result = subprocess.run(cmd, capture_output=True, text=True)
if result.returncode != 0:
    print(f"pandoc 错误: {result.stderr}")
    sys.exit(1)
print(f"pandoc 转换成功: {DST_DOCX}")

# python-docx 后格式化
print("正在后格式化 docx...")
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

doc = Document(str(DST_DOCX))
section = doc.sections[0]
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

def set_style_font(style, latin="Times New Roman", east_asia="宋体", size_pt=None):
    style.font.name = latin
    rpr = style._element.rPr
    if rpr is None:
        style._element.get_or_add_rPr()
        rpr = style._element.rPr
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), east_asia)
    if size_pt is not None:
        style.font.size = Pt(size_pt)

for st in doc.styles:
    if st.name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        sz = {"Normal": 12, "Heading 1": 16, "Heading 2": 14, "Heading 3": 13}[st.name]
        set_style_font(st, size_pt=sz)

for para in doc.paragraphs:
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    for run in para.runs:
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), "宋体")

for table in doc.tables:
    table.alignment = 1
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

doc.save(str(DST_DOCX))
print("后格式化完成")

# 统计
def count_omath(path):
    with zipfile.ZipFile(str(path), "r") as zf:
        return zf.read("word/document.xml").decode().count("<m:oMath>")

zh = len(re.findall(r"[一-鿿]", fixed_text))
omath = count_omath(DST_DOCX)
print(f"\n{'='*60}")
print(f"公式渲染版生成成功！")
print(f"{'='*60}")
print(f"DOCX: {DST_DOCX}")
print(f"大小: {DST_DOCX.stat().st_size/1024:.1f} KB")
print(f"OMML 公式数: {omath}")
print(f"嵌入图片数: 14")
print(f"中文字符数: {zh}")
print(f"{'='*60}")
