#!/usr/bin/env python3
"""
格式修复脚本 — 生成最终排版修正版（第二版）
========================================
读取当前 paper.md，将公式转为可读纯文本，修复表格和图题格式。
"""
import os, re, shutil, subprocess, sys, warnings, zipfile
from pathlib import Path

warnings.filterwarnings('ignore')
os.environ["PATH"] = "/usr/bin:" + os.environ.get("PATH", "")

BASE = Path("/root/ur10_ws/src/mud_experiment_recorder/data/analysis_results_20260513_005805")
PAPER_DIR = BASE / "paper"
FIG_DIR = BASE / "figures"
ASSET_DIR = PAPER_DIR / "assets"
SRC_MD = PAPER_DIR / "paper.md"
DST_MD = PAPER_DIR / "基于响应面法的UR10机械臂铲泥工艺参数多目标优化分析_最终排版修正版.md"
DST_DOCX = PAPER_DIR / "基于响应面法的UR10机械臂铲泥工艺参数多目标优化分析_最终排版修正版.docx"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
for f in sorted(FIG_DIR.glob("*.png")):
    shutil.copy2(f, ASSET_DIR / f.name)

# ============================================================
# 自定义 LaTeX → 纯文本 转换（支持嵌套括号）
# ============================================================

def _find_brace(s: str, start: int) -> int:
    """从 s[start]=='{' 找到匹配的 '}' 位置。"""
    if start >= len(s) or s[start] != '{':
        return -1
    d = 0
    for i in range(start, len(s)):
        if s[i] == '{':
            d += 1
        elif s[i] == '}':
            d -= 1
            if d == 0:
                return i
    return -1

def _extract_braced(s: str, start: int):
    """提取 s[start] 开始的括号内容（含{}），返回 (inner_text, end_pos)。"""
    if start >= len(s) or s[start] != '{':
        return None, start
    end = _find_brace(s, start)
    if end == -1:
        return s[start+1:], start + 1
    return s[start+1:end], end + 1

def latex_to_plain(text: str, is_inline: bool = False) -> str:
    """将 LaTeX 数学表达式转为可读纯文本（处理嵌套花括号）。"""
    out = []
    i = 0
    while i < len(text):
        c = text[i]
        # LaTeX 命令
        if c == '\\' and i + 1 < len(text):
            j = i + 1
            while j < len(text) and text[j].isalpha():
                j += 1
            cmd = text[i+1:j]
            if cmd in ('mathrm', 'text', 'mathbf'):
                inner, i_pos = _extract_braced(text, j)
                inner_plain = latex_to_plain(inner) if inner else ''
                out.append(inner_plain)
                i = i_pos
                continue
            elif cmd == 'frac':
                # 跳过空格到第一个 {
                k = j
                while k < len(text) and text[k] not in ('{', '}'):
                    k += 1
                num, k = _extract_braced(text, k)
                if num is None:
                    out.append('frac')
                    i = j
                    continue
                while k < len(text) and text[k] not in ('{', '}'):
                    k += 1
                den, k = _extract_braced(text, k)
                num_plain = latex_to_plain(num)
                den_plain = latex_to_plain(den)
                out.append(f'{num_plain}/{den_plain}')
                i = k
                continue
            elif cmd == 'sqrt':
                k = j
                while k < len(text) and text[k] not in ('{', '}'):
                    k += 1
                inner, k = _extract_braced(text, k)
                inner_plain = latex_to_plain(inner) if inner else ''
                out.append(f'sqrt({inner_plain})')
                i = k
                continue
            elif cmd == 'hat':
                k = j
                while k < len(text) and text[k] not in ('{', '}'):
                    k += 1
                inner, k = _extract_braced(text, k)
                out.append(f'{inner}̂' if inner else '')
                i = k
                continue
            elif cmd == 'dot':
                k = j
                while k < len(text) and text[k] not in ('{', '}'):
                    k += 1
                inner, k = _extract_braced(text, k)
                out.append(f'{inner}̇' if inner else '')
                i = k
                continue
            elif cmd == 'quad' or cmd == 'qquad':
                out.append('  ')
                i = j
                continue
            elif cmd in (',', ';', '!'):
                i = j
                continue
            # Greek 小写
            greek_map = {
                'alpha': 'α', 'beta': 'β', 'gamma': 'γ',
                'delta': 'δ', 'epsilon': 'ε', 'varepsilon': 'ε',
                'mu': 'μ', 'sigma': 'σ', 'tau': 'τ',
                'omega': 'ω', 'phi': 'φ', 'theta': 'θ',
                'eta': 'η', 'lambda': 'λ', 'pi': 'π',
            }
            if cmd in greek_map:
                out.append(greek_map[cmd])
                i = j
                continue
            # 符号
            sym_map = {
                'to': '→', 'rightarrow': '→', 'leftarrow': '←',
                'infty': '∞', 'partial': '∂', 'cdot': '·',
                'dots': '...', 'cdots': '...', 'sum': 'Σ',
                'int': '∫', 'max': 'max', 'min': 'min',
                'Delta': 'Δ', 'approx': '≈', 'simeq': '≃',
                'times': '×', 'pm': '±', 'mp': '∓',
                'ge': '≥', 'le': '≤', 'neq': '≠',
            }
            # \big, \bigl, \bigr, \Big 等尺寸命令 → 直接去除
            if cmd in ('big', 'Big', 'bigg', 'Bigg', 'bigl', 'Bigl',
                       'bigr', 'Bigr', 'biggl', 'Biggl', 'biggr', 'Biggr'):
                i = j
                continue
            if cmd in sym_map:
                out.append(sym_map[cmd])
                i = j
                continue
            # 单字符非字母命令（如 \, \; \! \\）
            if not cmd and len(text) > i + 1:
                single = text[i+1]
                if single in (',', ';', '!', ':', ' '):
                    i += 2
                    continue
                elif single == '\\':
                    i += 2
                    continue
            # 无法识别的命令，保留原名（去掉反斜线）
            if cmd:
                out.append(cmd)
                i = j
                continue
            # 孤立的 \ 字符
            out.append('\\')
            i += 1
        elif c == '{' or c == '}':
            # 去掉花括号（仅作分组用）
            i += 1
        elif c == '^':
            # 处理上标
            if i + 1 < len(text):
                if text[i+1] == '{':
                    inner, pos = _extract_braced(text, i+1)
                    if inner:
                        superscript = latex_to_plain(inner)
                        if superscript == '2':
                            out.append('²')
                        elif superscript == '3':
                            out.append('³')
                        elif superscript == '*':
                            out.append('*')
                        elif len(superscript) == 1:
                            out.append(superscript)
                        else:
                            out.append('^' + superscript)
                        i = pos
                        continue
                elif text[i+1] == '2':
                    out.append('²')
                    i += 2
                    continue
                elif text[i+1] == '3':
                    out.append('³')
                    i += 2
                    continue
            out.append('^')
            i += 1
        elif c == '_':
            # 处理下标（递归转换内部内容）
            if i + 1 < len(text) and text[i+1] == '{':
                inner, pos = _extract_braced(text, i+1)
                inner_plain = latex_to_plain(inner) if inner else ''
                out.append(f'_{inner_plain}')
                i = pos
                continue
            elif i + 1 < len(text):
                out.append(f'_{text[i+1]}')
                i += 2
                continue
            out.append('_')
            i += 1
        else:
            out.append(c)
            i += 1
    result = ''.join(out)
    # 清理
    result = result.replace('  ', ' ').strip()
    return result

# ============================================================
# Markdown 修复
# ============================================================

def fix_markdown(md_text: str) -> str:
    """综合格式修复。"""

    # 1. 摘要中 35°/35% 错误
    md_text = md_text.replace('对应 20°、35°、60 mm', '对应 20°、35%、60 mm')
    # 也修结论中类似问题
    md_text = re.sub(r'(参数为|角 )20°、35°、60 mm', r'\g<1>20°、35%、60 mm', md_text)

    # 2. 转换 $$...$$ 显示公式为代码块
    def _replace_display(m):
        content = m.group(1).strip()
        plain = latex_to_plain(content, is_inline=False)
        lines = [l for l in plain.split('\n') if l.strip()]
        if not lines:
            return ''
        code = '\n'.join(lines)
        return f'\n```\n{code}\n```\n'

    md_text = re.sub(r'\$\$(.*?)\$\$', _replace_display, md_text, flags=re.DOTALL)

    # 3. 转换 $...$ 内联公式（不处理已转义的 \$）
    def _replace_inline(m):
        content = m.group(1)
        plain = latex_to_plain(content, is_inline=True)
        return plain

    md_text = re.sub(r'(?<!\\)\$(.+?)(?<!\\)\$', _replace_inline, md_text)

    # 4. 清理残留的 $$、$
    md_text = md_text.replace('$$', '').replace('$', '')

    # 5. 模型拟合质量表头
    md_text = md_text.replace(
        '| 响应变量 | R² | R²_adj | RMSE | MAE |',
        '| 响应变量 | R² | Adjusted R² | RMSE | MAE |'
    )
    # 如果还残留 LaTeX 表头变体
    md_text = md_text.replace(
        '| 响应变量 | R² | R²_adj | RMSE | MAE |',
        '| 响应变量 | R² | Adjusted R² | RMSE | MAE |'
    )

    # 6. 清理多余空行（3+ → 2）
    md_text = re.sub(r'\n{4,}', '\n\n\n', md_text)

    return md_text


# ============================================================
# 执行
# ============================================================

src_text = SRC_MD.read_text(encoding='utf-8')
print(f"已读取: {SRC_MD} ({len(src_text)} chars)")

fixed_text = fix_markdown(src_text)
print(f"格式修正完成 ({len(fixed_text)} chars)")

DST_MD.write_text(fixed_text, encoding='utf-8')
print(f"已写入: {DST_MD}")

# pandoc 转换（不带 tex_math_dollars）
print("正在执行 pandoc 转换...")
cmd = [
    "pandoc", str(DST_MD),
    "-f", "markdown+grid_tables+pipe_tables",
    "-t", "docx", "-s",
    "--resource-path", str(PAPER_DIR),
    "-o", str(DST_DOCX),
]
result = subprocess.run(cmd, capture_output=True, text=True)
if result.returncode != 0:
    print(f"pandoc 错误: {result.stderr}")
    sys.exit(1)
print(f"pandoc 转换成功: {DST_DOCX}")

# python-docx 后格式化
print("正在后格式化 docx...")
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

doc = Document(str(DST_DOCX))
section = doc.sections[0]
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

def set_style_font(style, latin="Times New Roman", east_asia="宋体", size_pt=None):
    style.font.name = latin
    rpr = style._element.rPr
    if rpr is None:
        style._element.get_or_add_rPr()
        rpr = style._element.rPr
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), east_asia)
    if size_pt is not None:
        style.font.size = Pt(size_pt)

for st in doc.styles:
    if st.name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        sz = {"Normal": 12, "Heading 1": 16, "Heading 2": 14, "Heading 3": 13}[st.name]
        set_style_font(st, size_pt=sz)

for para in doc.paragraphs:
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    for run in para.runs:
        run.font.name = "Times New Roman"

for table in doc.tables:
    table.alignment = 1
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

doc.save(str(DST_DOCX))
print("后格式化完成")

# 统计
def count_omath(path):
    with zipfile.ZipFile(str(path), "r") as zf:
        return zf.read("word/document.xml").decode().count("<m:oMath>")

zh = len(re.findall(r"[一-鿿]", fixed_text))
print(f"\n{'='*60}")
print(f"最终排版修正版生成成功！")
print(f"{'='*60}")
print(f"DOCX: {DST_DOCX}")
print(f"大小: {DST_DOCX.stat().st_size/1024:.1f} KB")
print(f"OMML 公式数: {count_omath(DST_DOCX)}（应为 0）")
print(f"中文字符数: {zh}")
print(f"{'='*60}")
